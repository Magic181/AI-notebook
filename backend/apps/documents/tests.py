from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.test import TestCase, TransactionTestCase, override_settings
from docx import Document as DocxDocument
from rest_framework import status
from rest_framework.test import APIClient

from apps.notebooks.models import Notebook

from .chunking import chunk_blocks
from .models import Document, DocumentChunk, DocumentStatus
from .parsers import parse_file, parse_file_blocks
from .tasks import parse_document_task


class DocumentUploadTests(TransactionTestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.settings_override = override_settings(
            MEDIA_ROOT=Path(self.temp_dir.name),
            MAX_UPLOAD_SIZE_MB=20,
        )
        self.settings_override.enable()

        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username='upload-user',
            email='upload@example.com',
            password='test-password',
        )
        self.notebook = Notebook.objects.create(
            user=self.user,
            name='Upload notebook',
        )
        self.client = APIClient()
        self.client.force_authenticate(self.user)

    def tearDown(self):
        self.settings_override.disable()
        self.temp_dir.cleanup()

    def test_rejects_mixed_invalid_upload_before_saving_any_file(self):
        valid_file = self._file('valid.txt', b'valid content')
        invalid_file = self._file('slides.pptx', b'not supported')

        with patch('apps.documents.views.save_uploaded_file') as save_uploaded_file:
            response = self.client.post(
                f'/api/v1/notebooks/{self.notebook.id}/documents/',
                {'files': [valid_file, invalid_file]},
                format='multipart',
            )

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertFalse(save_uploaded_file.called)
        self.assertEqual(Document.objects.count(), 0)

    def test_valid_upload_creates_document_and_enqueues_parse_after_commit(self):
        uploaded_file = self._file('notes.md', b'# Notes\n\nBatch upload test.')

        with patch('apps.documents.views.enqueue_parse_task') as enqueue_parse_task:
            response = self.client.post(
                f'/api/v1/notebooks/{self.notebook.id}/documents/',
                {'files': [uploaded_file]},
                format='multipart',
            )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        document = Document.objects.get()
        self.assertEqual(document.name, 'notes.md')
        self.assertEqual(document.file_type, 'md')
        self.assertEqual(document.status, DocumentStatus.UPLOADING)
        self.assertTrue((Path(self.temp_dir.name) / document.file_path).exists())
        enqueue_parse_task.assert_called_once_with(document.id)

    @staticmethod
    def _file(name: str, content: bytes):
        from django.core.files.uploadedfile import SimpleUploadedFile

        return SimpleUploadedFile(name, content)


class DocxParsingTests(TestCase):
    def test_docx_parser_preserves_paragraph_table_order(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / 'report.docx'
            self._write_docx_with_table(path)

            blocks = parse_file_blocks(path, 'docx')
            text = parse_file(path, 'docx')

        self.assertEqual(
            [block['metadata']['source_type'] for block in blocks],
            ['paragraph', 'table', 'paragraph'],
        )
        self.assertIn('实验目标', blocks[0]['content'])
        self.assertIn('模块 | 状态', blocks[1]['content'])
        self.assertIn('行 2: 登录 | 通过', blocks[1]['content'])
        self.assertIn('实验结论', blocks[2]['content'])
        self.assertLess(text.index('实验目标'), text.index('[表格 1]'))
        self.assertLess(text.index('[表格 1]'), text.index('实验结论'))

    def test_chunk_blocks_keeps_table_metadata(self):
        chunks = chunk_blocks([
            {
                'content': '实验目标',
                'metadata': {'source_type': 'paragraph', 'block_index': 0},
            },
            {
                'content': '[表格 1]\n行 1: 模块 | 状态',
                'metadata': {
                    'source_type': 'table',
                    'block_index': 1,
                    'table_index': 1,
                    'row_count': 1,
                    'col_count': 2,
                },
            },
        ])

        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0]['metadata']['source_type'], 'paragraph')
        self.assertEqual(chunks[1]['metadata']['source_type'], 'table')
        self.assertEqual(chunks[1]['metadata']['table_index'], 1)

    @staticmethod
    def _write_docx_with_table(path: Path) -> None:
        doc = DocxDocument()
        doc.add_paragraph('实验目标：验证 DOCX 表格解析。')
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '模块'
        table.cell(0, 1).text = '状态'
        table.cell(1, 0).text = '登录'
        table.cell(1, 1).text = '通过'
        doc.add_paragraph('实验结论：表格内容已被读取。')
        doc.save(path)


class ParseDocumentTaskTests(TransactionTestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.settings_override = override_settings(MEDIA_ROOT=Path(self.temp_dir.name))
        self.settings_override.enable()

        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username='parse-user',
            email='parse@example.com',
            password='test-password',
        )
        self.notebook = Notebook.objects.create(user=self.user, name='Parse notebook')

    def tearDown(self):
        self.settings_override.disable()
        self.temp_dir.cleanup()

    def test_parse_task_stores_docx_table_chunks_with_metadata(self):
        relative_path = Path('files') / str(self.user.id) / str(self.notebook.id) / 'report.docx'
        file_path = Path(self.temp_dir.name) / relative_path
        file_path.parent.mkdir(parents=True, exist_ok=True)
        DocxParsingTests._write_docx_with_table(file_path)

        document = Document.objects.create(
            notebook=self.notebook,
            name='report.docx',
            file_path=str(relative_path),
            file_size=file_path.stat().st_size,
            file_type='docx',
            status=DocumentStatus.UPLOADING,
        )

        parse_document_task(document.id)

        document.refresh_from_db()
        self.assertEqual(document.status, DocumentStatus.COMPLETED)
        self.assertGreaterEqual(document.chunk_count, 2)
        table_chunk = DocumentChunk.objects.get(
            document=document,
            metadata__source_type='table',
        )
        self.assertIn('模块 | 状态', table_chunk.content)
        self.assertEqual(table_chunk.metadata['table_index'], 1)
